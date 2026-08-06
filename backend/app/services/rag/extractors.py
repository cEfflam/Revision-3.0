"""
Extraction de texte brut depuis un fichier importé.

Objectif : produire du Markdown. Pas par coquetterie — les titres `#` sont le
signal dont le découpeur se sert pour couper aux frontières sémantiques plutôt
qu'au milieu d'une phrase. Un chunk qui commence par « ...et donc la jointure »
est inexploitable pour le RAG.

Formats couverts : PDF, DOCX, Markdown, texte brut.
Les images/scans (OCR) arrivent en phase 2 : ça demande Tesseract dans l'image
Docker, ce qui double sa taille. On attend d'en avoir vraiment besoin.
"""

from __future__ import annotations

import io
import logging
import re
from dataclasses import dataclass
from pathlib import Path

logger = logging.getLogger(__name__)

MARKDOWN_EXTENSIONS = {".md", ".markdown", ".mdown"}
TEXT_EXTENSIONS = {".txt", ".text", ".csv", ".log"}
PDF_EXTENSIONS = {".pdf"}
DOCX_EXTENSIONS = {".docx"}
OCR_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp", ".tiff", ".bmp"}

SUPPORTED_EXTENSIONS = (
    MARKDOWN_EXTENSIONS | TEXT_EXTENSIONS | PDF_EXTENSIONS | DOCX_EXTENSIONS
)


class UnsupportedDocument(ValueError):
    """Format non pris en charge — message destiné à l'utilisateur final."""


@dataclass(slots=True)
class ExtractedDocument:
    text: str
    page_count: int = 0
    detected_title: str = ""


def extract_text(data: bytes, filename: str) -> ExtractedDocument:
    """Point d'entrée unique : aiguille selon l'extension du fichier."""
    suffix = Path(filename).suffix.lower()

    if suffix in PDF_EXTENSIONS:
        return _extract_pdf(data, filename)
    if suffix in DOCX_EXTENSIONS:
        return _extract_docx(data, filename)
    if suffix in MARKDOWN_EXTENSIONS or suffix in TEXT_EXTENSIONS:
        return _extract_plain(data, filename)
    if suffix in OCR_EXTENSIONS:
        raise UnsupportedDocument(
            "Les images et scans nécessitent l'OCR, prévu en phase 2. "
            "En attendant, exporte ton document en PDF texte ou en Markdown."
        )
    raise UnsupportedDocument(
        f"Format « {suffix or 'inconnu'} » non pris en charge. "
        f"Formats acceptés : {', '.join(sorted(SUPPORTED_EXTENSIONS))}."
    )


def _decode(data: bytes) -> str:
    """UTF-8 d'abord, puis les encodages Windows classiques."""
    for encoding in ("utf-8", "utf-8-sig", "cp1252", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def _extract_plain(data: bytes, filename: str) -> ExtractedDocument:
    text = _clean(_decode(data))
    # Un .md est censé porter sa propre syntaxe : on n'y touche pas, au risque
    # sinon de créer des titres concurrents de ceux de l'auteur. Un .txt, lui,
    # n'a aucune structure — même traitement que pour un PDF.
    if Path(filename).suffix.lower() in TEXT_EXTENSIONS:
        text = _promote_headings(text)
    return ExtractedDocument(
        text=text, detected_title=_first_heading(text) or Path(filename).stem
    )


def _extract_pdf(data: bytes, filename: str) -> ExtractedDocument:
    from pypdf import PdfReader

    try:
        reader = PdfReader(io.BytesIO(data))
    except Exception as exc:  # pypdf lève des exceptions très variées
        raise UnsupportedDocument(f"PDF illisible ou corrompu : {exc}") from exc

    if reader.is_encrypted:
        # Beaucoup de PDF « protégés » le sont avec un mot de passe vide.
        try:
            reader.decrypt("")
        except Exception as exc:
            raise UnsupportedDocument(
                "Ce PDF est protégé par mot de passe."
            ) from exc

    parts: list[str] = []
    for index, page in enumerate(reader.pages, start=1):
        try:
            content = page.extract_text() or ""
        except Exception:
            logger.warning("Page %s illisible dans %s", index, filename)
            content = ""
        content = _promote_headings(_clean(content))
        if content:
            # La page n'est plus un titre mais un simple repère : les vrais
            # titres viennent d'être reconstruits par _promote_headings, et
            # deux niveaux de « Page N » écraseraient cette hiérarchie.
            parts.append(f"<!-- page {index} -->\n{content}")

    if not parts:
        raise UnsupportedDocument(
            "Aucun texte extractible : ce PDF est probablement un scan. "
            "L'OCR est prévu en phase 2."
        )

    text = "\n\n".join(parts)
    title = _pdf_title(reader) or Path(filename).stem
    return ExtractedDocument(
        text=text, page_count=len(reader.pages), detected_title=title
    )


def _pdf_title(reader) -> str:
    try:
        raw = (reader.metadata or {}).get("/Title", "")
        return str(raw).strip()
    except Exception:
        return ""


def _extract_docx(data: bytes, filename: str) -> ExtractedDocument:
    import docx  # python-docx

    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as exc:
        raise UnsupportedDocument(f"DOCX illisible : {exc}") from exc

    lines: list[str] = []
    for paragraph in document.paragraphs:
        content = paragraph.text.strip()
        if not content:
            continue
        # Les styles Word « Heading 1/2/3 » sont convertis en titres Markdown :
        # c'est la structure du cours qu'on récupère, gratuitement.
        style = (paragraph.style.name or "").lower()
        if style.startswith(("heading", "titre")):
            digits = re.findall(r"\d+", style)
            level = min(int(digits[0]), 6) if digits else 1
            lines.append(f"{'#' * level} {content}")
        else:
            lines.append(content)

    # Les tableaux sont fréquents dans les cours (comparatifs, barèmes).
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                lines.append(" | ".join(cells))

    text = _clean("\n\n".join(lines))
    if not text:
        raise UnsupportedDocument("Ce document Word ne contient aucun texte.")

    return ExtractedDocument(
        text=text, detected_title=_first_heading(text) or Path(filename).stem
    )


# ── Reconstruction de la structure d'un PDF ──────────────────────────────
# Un PDF ne contient aucune notion de « titre » : pypdf en sort un flot de
# lignes. Sans traitement, chaque fragment n'a pour tout contexte que
# « Page 24 » — inutilisable pour citer une source ou cibler une recherche.
# Ces motifs retrouvent la hiérarchie que l'œil humain lit d'un coup.

#: « Chapitre 5 : L'organisation de l'activité » → titre de niveau 2
_CHAPTER_RE = re.compile(
    r"^(?:chapitre|partie|module|dossier)\s+[\dIVX]+\s*[:.\-–]?\s*(.{3,90})$",
    flags=re.IGNORECASE,
)
#: « 9. Les différents processus de l'entreprise : » → titre de niveau 3
_NUMBERED_RE = re.compile(r"^\d{1,2}[.)]\s+(.{3,90}?)\s*:\s*$")
#: Ligne de sommaire « Chapitre 5 ...................... 60 » : à supprimer.
#: Réindexer un sommaire pollue la recherche avec des entrées sans contenu.
_TOC_RE = re.compile(r"\.{5,}\s*\d{1,4}\s*$")
#: Ligne ne contenant qu'un numéro de page.
_PAGE_NUMBER_RE = re.compile(r"^\d{1,4}$")


def _promote_headings(text: str) -> str:
    """Transforme les repères visuels d'un PDF en titres Markdown."""
    lines: list[str] = []
    for raw in text.split("\n"):
        line = raw.strip()
        if not line or _PAGE_NUMBER_RE.match(line) or _TOC_RE.search(line):
            continue

        chapter = _CHAPTER_RE.match(line)
        if chapter:
            lines.append(f"## {chapter.group(1).strip(' :')}")
            continue

        numbered = _NUMBERED_RE.match(line)
        if numbered:
            lines.append(f"### {numbered.group(1).strip()}")
            continue

        lines.append(line)
    return "\n".join(lines)


def _clean(text: str) -> str:
    """Normalise les sauts de ligne et supprime les espaces parasites."""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = text.replace(" ", " ")          # espace insécable
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)      # max 1 ligne vide
    return "\n".join(line.strip() for line in text.split("\n")).strip()


def _first_heading(text: str) -> str:
    match = re.search(r"^#{1,3}\s+(.+)$", text, flags=re.MULTILINE)
    return match.group(1).strip() if match else ""
